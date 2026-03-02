
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_style(doc):
    # 设置正文字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 设置标题样式
    for i in range(1, 4):
        style_name = f'Heading {i}'
        if style_name not in doc.styles:
            continue
        style = doc.styles[style_name]
        style.font.name = 'Arial'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.color.rgb = RGBColor(0, 0, 0)
        if i == 1:
            style.font.size = Pt(16)
            style.paragraph_format.space_before = Pt(24)
            style.paragraph_format.space_after = Pt(18)
        elif i == 2:
            style.font.size = Pt(14)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(12)
        else:
            style.font.size = Pt(13)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def generate_srs():
    doc = Document()
    set_style(doc)

    # --- 封面 ---
    for _ in range(5):
        doc.add_paragraph()
    
    title = doc.add_paragraph('软件需求规格说明书')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.style.font.size = Pt(26)
    title.style.font.bold = True

    subtitle = doc.add_paragraph('(GB/T 8567-2006)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.style.font.size = Pt(14)

    for _ in range(5):
        doc.add_paragraph()

    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run('项目名称：童康源家庭教育支持系统\n').bold = True
    info_para.add_run('文件状态：正式发布\n')
    info_para.add_run('版本号：V1.0\n')
    info_para.add_run('生成日期：2026-01-06')

    doc.add_page_break()

    # --- 目录 (Word自动生成目录比较复杂，这里仅占位，实际需在Word中更新) ---
    # doc.add_paragraph('目  录').alignment = WD_ALIGN_PARAGRAPH.CENTER
    # doc.add_page_break()

    # --- 1. 引言 ---
    add_heading(doc, '1. 引言', level=1)
    
    add_heading(doc, '1.1 标识', level=2)
    doc.add_paragraph('本软件名为“童康源家庭教育支持系统”（TongYangYuan），是一个集成了移动端应用（Android）、Web前端、后端服务（Spring Boot）及AI智能分析服务（YOLOv8）的综合性家庭教育支持平台。')

    add_heading(doc, '1.2 系统概述', level=2)
    doc.add_paragraph('童康源旨在为家庭提供全链路的教育支持与儿童心理健康服务。系统连接家长与专业咨询师，提供在线咨询、视频通话、成长档案管理、社区互动及AI行为分析等功能。通过数字化手段，解决家庭教育资源获取难、沟通不畅及缺乏科学指导的问题。')

    add_heading(doc, '1.3 文档概述', level=2)
    doc.add_paragraph('本文档依据GB/T 8567-2006标准编写，详细描述了童康源系统的功能需求、接口需求、性能需求及其他非功能需求。本文档面向项目开发人员、测试人员及系统维护人员。')

    # --- 2. 总体描述 ---
    add_heading(doc, '2. 总体描述', level=1)

    add_heading(doc, '2.1 用户特征', level=2)
    doc.add_paragraph('本系统的主要用户包括：')
    doc.add_paragraph('1. 家长用户：关注孩子成长，需要获取育儿知识、预约专家咨询、记录孩子成长历程。')
    doc.add_paragraph('2. 咨询师/导师：提供专业咨询服务，管理预约，进行视频/图文辅导。')
    doc.add_paragraph('3. 系统管理员：负责用户管理、内容审核、系统维护。')

    add_heading(doc, '2.2 运行环境', level=2)
    doc.add_paragraph('1. 客户端：Android 8.0及以上版本操作系统。')
    doc.add_paragraph('2. 服务端：支持Java运行环境（JDK 17+）的服务器，及Python 3.8+运行环境（用于AI服务）。')
    doc.add_paragraph('3. 数据库：MySQL 8.0+。')
    doc.add_paragraph('4. 网络环境：支持4G/5G/Wi-Fi网络。')

    # --- 3. 具体需求 ---
    add_heading(doc, '3. 具体需求', level=1)

    add_heading(doc, '3.1 功能需求', level=2)

    # 3.1.1 用户认证与管理
    add_heading(doc, '3.1.1 用户认证与管理', level=3)
    p = doc.add_paragraph()
    p.add_run('功能描述：').bold = True
    p.add_run('提供用户注册、登录、密码管理及个人信息维护功能。')
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '功能项'
    table.rows[0].cells[1].text = '描述'
    table.rows[0].cells[2].text = '优先级'
    
    table.rows[1].cells[0].text = '注册/登录'
    table.rows[1].cells[1].text = '支持手机号/用户名注册，使用Token机制维持会话。'
    table.rows[1].cells[2].text = '高'

    table.rows[2].cells[0].text = '首次引导'
    table.rows[2].cells[1].text = '新用户首次登录后需完善孩子基本信息（姓名、性别、生日、兴趣等）。'
    table.rows[2].cells[2].text = '中'

    table.rows[3].cells[0].text = '会员服务'
    table.rows[3].cells[1].text = '支持查看会员权益及充值操作（模拟支付流程）。'
    table.rows[3].cells[2].text = '中'

    # 3.1.2 咨询服务
    add_heading(doc, '3.1.2 咨询服务', level=3)
    p = doc.add_paragraph()
    p.add_run('功能描述：').bold = True
    p.add_run('核心业务功能，连接家长与咨询师。')
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = '功能项'
    table.rows[0].cells[1].text = '描述'
    table.rows[0].cells[2].text = '优先级'
    
    table.rows[1].cells[0].text = '导师列表'
    table.rows[1].cells[1].text = '展示咨询师信息（头像、资质、评分、擅长领域），支持点击查看详情。'
    table.rows[1].cells[2].text = '高'

    table.rows[2].cells[0].text = '预约咨询'
    table.rows[2].cells[1].text = '选择日期和时间段进行预约，生成咨询订单。'
    table.rows[2].cells[2].text = '高'

    table.rows[3].cells[0].text = '即时通讯'
    table.rows[3].cells[1].text = '基于WebSocket (Stomp) 实现文字、图片、语音消息发送；集成WebRTC实现视频通话。'
    table.rows[3].cells[2].text = '高'

    # 3.1.3 社区广场
    add_heading(doc, '3.1.3 社区广场', level=3)
    doc.add_paragraph('用户可在“社交广场”浏览其他家长或导师发布的动态。')
    doc.add_paragraph('1. 动态展示：瀑布流或列表形式展示图文动态。')
    doc.add_paragraph('2. 互动功能：支持点赞（含动画效果）、评论。')
    doc.add_paragraph('3. 发布动态：支持上传图片（相册/拍照）、输入文字发布新动态。')

    # 3.1.4 AI 智能识别
    add_heading(doc, '3.1.4 AI 智能识别', level=3)
    doc.add_paragraph('集成 YOLOv8 计算机视觉模型，提供智能分析服务。')
    doc.add_paragraph('1. 图像采集：支持通过 App 调用摄像头拍照或从相册选择图片。')
    doc.add_paragraph('2. 智能分析：图片上传至 Python 后端，识别场景中的物体或行为特征（如儿童行为分析、环境评估）。')
    doc.add_paragraph('3. 结果反馈：在前端实时绘制识别框（Bounding Box），显示类别标签及置信度。')
    doc.add_paragraph('4. 仪表盘入口：在首页仪表盘提供快捷入口。')

    add_heading(doc, '3.2 外部接口需求', level=2)
    
    add_heading(doc, '3.2.1 用户接口', level=3)
    doc.add_paragraph('采用混合开发模式（Hybrid App），主要界面通过 WebView 加载 HTML5 页面，底部导航栏及部分核心功能（如视频通话、媒体采集）采用 Android 原生实现。')
    
    add_heading(doc, '3.2.2 硬件接口', level=3)
    doc.add_paragraph('1. 摄像头：用于拍照、视频通话、AI识别。')
    doc.add_paragraph('2. 麦克风：用于语音消息、视频通话。')
    doc.add_paragraph('3. 存储：用于保存聊天记录、缓存图片。')

    add_heading(doc, '3.2.3 软件接口', level=3)
    doc.add_paragraph('1. 后端 API：基于 HTTP/HTTPS 协议的 RESTful 接口。')
    doc.add_paragraph('2. 消息推送：WebSocket 连接，用于实时消息推送。')
    doc.add_paragraph('3. AI 服务接口：Python Flask 服务提供的 /predict 接口。')

    add_heading(doc, '3.3 性能需求', level=2)
    doc.add_paragraph('1. 响应时间：页面加载时间不超过 2 秒；API 接口响应时间不超过 500ms。')
    doc.add_paragraph('2. AI 识别效率：单张图片识别处理时间不超过 3 秒。')
    doc.add_paragraph('3. 并发能力：支持至少 100 用户同时在线操作。')

    add_heading(doc, '3.4 安全性需求', level=2)
    doc.add_paragraph('1. 数据传输：敏感数据（如密码、Token）传输需加密。')
    doc.add_paragraph('2. 权限控制：严格控制文件访问、摄像头及录音权限，遵循 Android 权限模型。')
    doc.add_paragraph('3. 输入验证：对所有用户输入进行验证，防止 SQL 注入及 XSS 攻击。')

    # --- 4. 附录 ---
    add_heading(doc, '4. 附录', level=1)
    
    add_heading(doc, '附录 A：数据字典', level=2)
    doc.add_paragraph('本附录详细列出了系统核心实体的数据结构定义。')

    def add_data_table(doc, title, rows_data):
        p = doc.add_paragraph(title)
        p.style = 'Heading 3'
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '字段名'
        hdr_cells[1].text = '数据类型'
        hdr_cells[2].text = '长度/范围'
        hdr_cells[3].text = '是否必填'
        hdr_cells[4].text = '描述'
        
        for field, dtype, length, required, desc in rows_data:
            row_cells = table.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = dtype
            row_cells[2].text = length
            row_cells[3].text = required
            row_cells[4].text = desc
        doc.add_paragraph() # Add spacing

    # 1. 用户表 (users)
    user_data = [
        ('id', 'Long', '20', '是', '主键 ID'),
        ('phone', 'String', '20', '是', '手机号（唯一）'),
        ('password', 'String', '255', '是', '加密密码'),
        ('user_type', 'Enum', '-', '是', '用户类型 (CONSULTANT/PARENT)'),
        ('nickname', 'String', '50', '否', '昵称'),
        ('avatar_url', 'String', '500', '否', '头像 URL'),
        ('status', 'Enum', '-', '是', '状态 (ACTIVE/BANNED)'),
        ('last_login_at', 'DateTime', '-', '否', '最后登录时间')
    ]
    add_data_table(doc, 'A.1 用户表 (User / users)', user_data)

    # 2. 儿童档案表 (children)
    child_data = [
        ('id', 'Long', '20', '是', '主键 ID'),
        ('parent_user_id', 'Long', '20', '是', '关联家长用户 ID'),
        ('name', 'String', '50', '是', '真实姓名'),
        ('gender', 'Enum', '-', '否', '性别 (MALE/FEMALE)'),
        ('birth_date', 'Date', '-', '否', '出生日期'),
        ('school', 'String', '200', '否', '就读学校'),
        ('interests', 'Text', '-', '否', '兴趣爱好'),
        ('health_status', 'Text', '-', '否', '健康状况'),
        ('medical_history', 'Text', '-', '否', '病史')
    ]
    add_data_table(doc, 'A.2 儿童档案表 (Child / children)', child_data)

    # 3. 咨询师表 (consultants)
    consultant_data = [
        ('id', 'Long', '20', '是', '主键 ID'),
        ('user_id', 'Long', '20', '是', '关联用户 ID'),
        ('name', 'String', '50', '是', '咨询师姓名'),
        ('title', 'String', '100', '否', '职称/头衔'),
        ('specialty', 'Text', '-', '否', '擅长领域'),
        ('rating', 'Decimal', '3,2', '否', '评分 (0.00-5.00)'),
        ('served_count', 'Int', '-', '否', '服务次数'),
        ('is_available', 'Boolean', '-', '否', '是否可预约')
    ]
    add_data_table(doc, 'A.3 咨询师表 (Consultant / consultants)', consultant_data)

    # 4. 预约表 (appointments)
    appt_data = [
        ('id', 'Long', '20', '是', '主键 ID'),
        ('appointment_no', 'String', '50', '是', '预约编号'),
        ('consultant_id', 'Long', '20', '是', '咨询师 ID'),
        ('parent_user_id', 'Long', '20', '是', '家长 ID'),
        ('child_id', 'Long', '20', '否', '关联儿童 ID'),
        ('appointment_date', 'Date', '-', '是', '预约日期'),
        ('time_slot', 'String', '50', '是', '时间段 (如 10:00-11:00)'),
        ('status', 'Enum', '-', '是', '状态 (PENDING/COMPLETED/CANCELLED)'),
        ('description', 'Text', '-', '否', '咨询描述')
    ]
    add_data_table(doc, 'A.4 预约表 (Appointment / appointments)', appt_data)

    # 5. 社区动态表 (posts)
    post_data = [
        ('id', 'Long', '20', '是', '主键 ID'),
        ('author_user_id', 'Long', '20', '是', '发布者 ID'),
        ('title', 'Text', '-', '是', '标题/主要内容'),
        ('content', 'Text', '-', '否', '详细内容'),
        ('image_url', 'String', '500', '否', '图片链接'),
        ('like_count', 'Int', '-', '是', '点赞数'),
        ('created_at', 'DateTime', '-', '是', '发布时间')
    ]
    add_data_table(doc, 'A.5 社区动态表 (Post / posts)', post_data)

    add_heading(doc, '附录 B：缩略语清单', level=2)
    doc.add_paragraph('   - SRS: Software Requirements Specification')
    doc.add_paragraph('   - API: Application Programming Interface')
    doc.add_paragraph('   - YOLO: You Only Look Once (Object Detection Model)')
    doc.add_paragraph('   - JSON: JavaScript Object Notation')

    # 保存文件
    file_name = '童康源_软件需求规格说明书_v2.docx'
    try:
        doc.save(file_name)
        print(f"文档已生成: {os.path.abspath(file_name)}")
    except PermissionError:
        print(f"无法保存文件 '{file_name}'，请确保文件未被打开。尝试保存为 _v3...")
        file_name = '童康源_软件需求规格说明书_v3.docx'
        doc.save(file_name)
        print(f"文档已生成: {os.path.abspath(file_name)}")

if __name__ == '__main__':
    generate_srs()
